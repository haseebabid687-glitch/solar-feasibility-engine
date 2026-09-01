# ============================================================
# M2P Consulting – Airport Solar Feasibility Tool
# Version 3.1  |  31 August 2026
#
# v3.0 adds:
#   · Calculation methodology expanders (full traceability)
#   · Multi-sheet Excel export (Summary / Monthly / Cash flow)
#   · Word report generation (.docx)
#
# v3.1 adds:
#   · All areas in square metres (m2) rather than hectares
#   · Gate 7 — spatial waterfall chart and interactive site map
# ============================================================
import io
import math
from datetime import date

import streamlit as st
import pandas as pd
import numpy_financial as npf
import plotly.graph_objects as go

# folium is optional at runtime — the app must still load if it is missing
try:
    import folium
    from streamlit_folium import st_folium
    FOLIUM_AVAILABLE = True
except ImportError:
    FOLIUM_AVAILABLE = False

# folium / streamlit-folium are optional at runtime
try:
    import folium
    from streamlit_folium import st_folium
    FOLIUM_AVAILABLE = True
except ImportError:
    FOLIUM_AVAILABLE = False

# python-docx is optional at runtime — the app must still load if it is missing
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor as DocxRGB
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ── BRAND COLOURS ──────────────────────────────────────────
M2P_NAVY   = "#063443"
M2P_TEAL   = "#184A41"
M2P_BLUE   = "#2581C4"
M2P_GREEN  = "#008A65"
M2P_AMBER  = "#BD7119"
M2P_LIGHT  = "#F2F2F2"
RED        = "#C0392B"
AMBER_WARN = "#E67E22"
GREEN_OK   = "#27AE60"

# ── PAGE CONFIG ────────────────────────────────────────────
st.set_page_config(
    page_title="M2P Solar Feasibility Tool",
    page_icon="☀️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── GLOBAL CSS ─────────────────────────────────────────────
st.markdown(f"""
<style>
    /* Base */
    html, body, [class*="css"] {{
        font-family: 'Century Gothic', 'Trebuchet MS', sans-serif;
        background-color: #FAFAFA;
    }}
    /* Header bar */
    .m2p-header {{
        background: linear-gradient(135deg, {M2P_NAVY} 0%, {M2P_TEAL} 100%);
        padding: 1.4rem 2rem;
        border-radius: 0 0 8px 8px;
        margin-bottom: 1.5rem;
    }}
    .m2p-header h1 {{
        color: white;
        font-size: 1.6rem;
        font-weight: 700;
        margin: 0;
        letter-spacing: 0.3px;
    }}
    .m2p-header p {{
        color: rgba(255,255,255,0.72);
        font-size: 0.85rem;
        margin: 0.25rem 0 0 0;
    }}
    /* Gate cards */
    .gate-card {{
        background: white;
        border: 1px solid #E8E8E8;
        border-left: 4px solid {M2P_BLUE};
        border-radius: 6px;
        padding: 1.2rem 1.4rem;
        margin-bottom: 1rem;
        box-shadow: 0 1px 3px rgba(0,0,0,0.05);
    }}
    .gate-card h3 {{
        color: {M2P_NAVY};
        font-size: 1rem;
        font-weight: 700;
        margin: 0 0 0.6rem 0;
    }}
    .gate-card.warn {{ border-left-color: {AMBER_WARN}; }}
    .gate-card.pass {{ border-left-color: {GREEN_OK}; }}
    .gate-card.fail {{ border-left-color: {RED}; }}
    /* KPI tiles */
    .kpi-row {{ display: flex; gap: 1rem; margin: 1rem 0; flex-wrap: wrap; }}
    .kpi-tile {{
        flex: 1; min-width: 160px;
        background: {M2P_NAVY};
        border-radius: 8px;
        padding: 1.1rem 1.3rem;
        text-align: center;
    }}
    .kpi-tile .val {{
        font-size: 1.9rem;
        font-weight: 700;
        color: white;
        line-height: 1.1;
    }}
    .kpi-tile .lbl {{
        font-size: 0.72rem;
        color: rgba(255,255,255,0.65);
        margin-top: 0.3rem;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }}
    .kpi-tile.amber {{ background: {M2P_AMBER}; }}
    .kpi-tile.teal  {{ background: {M2P_TEAL}; }}
    .kpi-tile.blue  {{ background: {M2P_BLUE}; }}
    .kpi-tile.green {{ background: {M2P_GREEN}; }}
    /* Status badges */
    .badge {{
        display: inline-block;
        padding: 0.18rem 0.7rem;
        border-radius: 20px;
        font-size: 0.75rem;
        font-weight: 700;
        letter-spacing: 0.4px;
    }}
    .badge-pass {{ background: #D5F5E3; color: #1E8449; }}
    .badge-warn {{ background: #FDEBD0; color: #A04000; }}
    .badge-fail {{ background: #FADBD8; color: #922B21; }}
    .badge-info {{ background: #D6EAF8; color: #1A5276; }}
    /* Section labels */
    .section-label {{
        font-size: 0.7rem;
        font-weight: 700;
        color: {M2P_BLUE};
        text-transform: uppercase;
        letter-spacing: 1px;
        margin-bottom: 0.4rem;
    }}
    /* Divider */
    .m2p-divider {{
        border: none;
        border-top: 2px solid {M2P_LIGHT};
        margin: 1.5rem 0;
    }}
    /* Sidebar — base */
    [data-testid="stSidebar"] {{
        background: {M2P_NAVY};
    }}
    [data-testid="stSidebar"] * {{ color: rgba(255,255,255,0.88); }}

    /* Sidebar labels */
    [data-testid="stSidebar"] label,
    [data-testid="stSidebar"] .stSelectbox label,
    [data-testid="stSidebar"] .stNumberInput label,
    [data-testid="stSidebar"] .stSlider label,
    [data-testid="stSidebar"] .stTextInput label {{
        color: rgba(255,255,255,0.70) !important;
        font-size: 0.78rem;
    }}

    /* Sidebar input boxes — white background, dark text so values are readable */
    [data-testid="stSidebar"] input,
    [data-testid="stSidebar"] textarea,
    [data-testid="stSidebar"] .stTextInput input,
    [data-testid="stSidebar"] .stNumberInput input {{
        background-color: #FFFFFF !important;
        color: {M2P_NAVY} !important;
        border: 1px solid rgba(255,255,255,0.25) !important;
        border-radius: 4px !important;
    }}

    /* Sidebar select boxes */
    [data-testid="stSidebar"] .stSelectbox div[data-baseweb="select"] > div,
    [data-testid="stSidebar"] select {{
        background-color: #FFFFFF !important;
        color: {M2P_NAVY} !important;
        border-radius: 4px !important;
    }}

    /* Sidebar slider track and thumb — keep visible on dark bg */
    [data-testid="stSidebar"] .stSlider [data-testid="stTickBar"] {{
        color: rgba(255,255,255,0.50) !important;
    }}

    /* Sidebar markdown / headings */
    [data-testid="stSidebar"] h1,
    [data-testid="stSidebar"] h2,
    [data-testid="stSidebar"] h3,
    [data-testid="stSidebar"] p,
    [data-testid="stSidebar"] span:not([class]) {{
        color: rgba(255,255,255,0.88) !important;
    }}
    /* Disclaimer banner */
    .disclaimer {{
        background: #EBF5FB;
        border: 1px solid #AED6F1;
        border-radius: 6px;
        padding: 0.8rem 1.2rem;
        font-size: 0.78rem;
        color: #1A5276;
        margin-top: 2rem;
    }}
</style>
""", unsafe_allow_html=True)

# ── HEADER ─────────────────────────────────────────────────
st.markdown("""
<div class="m2p-header">
  <h1>☀️ Airport Solar Feasibility Tool</h1>
  <p>M2P Consulting · Airport Planning & Design · Internal Use Only · v3.0</p>
</div>
""", unsafe_allow_html=True)

# ── SIDEBAR ────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### Project Setup")
    project_name   = st.text_input("Project name", value="Airport Solar Assessment")
    airport_type   = st.selectbox("Airport development type", ["Brownfield", "Greenfield"])
    airport_region = st.selectbox("Region (for solar resource)", [
        "Europe – North (UK, Netherlands, Norway) · 2.4 h",
        "Europe – Central (Germany, France) · 2.9 h",
        "Europe – South (Spain, Italy, Greece) · 4.2 h",
        "Middle East (UAE, Qatar, Saudi) · 5.8 h",
        "South / SE Asia (Malaysia, India) · 4.6 h",
        "North America – East · 3.5 h",
        "North America – West (California) · 5.0 h",
        "Custom",
    ])
    custom_psh = None
    if "Custom" in airport_region:
        custom_psh = st.number_input("Peak sun hours (h/day)", 1.5, 7.0, 4.0, 0.1)

    psh_lookup = {
        "Europe – North": 2.4, "Europe – Central": 2.9, "Europe – South": 4.2,
        "Middle East": 5.8, "South / SE Asia": 4.6,
        "North America – East": 3.5, "North America – West": 5.0,
    }
    peak_sun = custom_psh if custom_psh else next(
        (v for k, v in psh_lookup.items() if k in airport_region), 3.5
    )

    st.markdown("---")
    st.markdown("### Market Assumptions")
    capex_mw       = st.number_input("Solar CAPEX (€ / MW)", 500000, 2000000, 800000, 50000,
                                      help="European brownfield rooftop ~€900k; Middle East ground-mount ~€650k")
    ppa_price      = st.number_input("PPA / avoided cost (€ / kWh)", 0.05, 0.25, 0.09, 0.01)
    opex_pct       = st.slider("OPEX (% of CAPEX / yr)", 0.5, 3.0, 1.5, 0.1) / 100
    ppa_escalation = st.slider("Annual PPA escalation (%)", 0.0, 3.0, 1.5, 0.5) / 100
    opex_inflation = st.slider("Annual OPEX inflation (%)", 0.0, 5.0, 2.5, 0.5) / 100

    st.markdown("---")
    st.markdown("### Thermal & Module")
    st.caption("Hot-climate derating. Defaults reflect Gulf summer conditions.")
    amb_temp_peak  = st.slider("Peak ambient temperature (°C)", 20, 55, 45, 1,
                                help="Summer design temperature. Doha/Riyadh peak ~45–50°C; Frankfurt ~30°C")
    amb_temp_swing = st.slider("Annual temperature swing (°C)", 5, 35, 25, 1,
                                help="Difference between summer peak and winter minimum monthly mean")
    temp_coeff     = st.slider("Temperature coefficient (%/°C)", -0.50, -0.20, -0.35, 0.01,
                                help="Power loss per °C above 25°C STC. Modern mono-PERC ≈ −0.35%/°C") / 100
    noct           = st.slider("NOCT (°C)", 40, 50, 45, 1,
                                help="Nominal Operating Cell Temperature from the module datasheet. Standard ≈ 45°C")

    bifacial       = st.checkbox("Bifacial modules", value=False,
                                  help="Rear-side generation. Requires clearance beneath the array — minimal benefit on flush-mounted roofs")
    albedo         = 0.25
    bifaciality    = 0.75
    if bifacial:
        albedo      = st.slider("Surface albedo", 0.10, 0.85, 0.30, 0.05,
                                 help="Desert sand 0.25–0.40 · White membrane roof 0.60–0.80 · Gravel 0.15–0.20 · Concrete 0.25–0.35")
        bifaciality = st.slider("Bifaciality factor", 0.60, 0.90, 0.75, 0.05,
                                 help="Rear-side efficiency relative to front, from the module datasheet")

    st.markdown("---")
    st.markdown("### Finance")
    debt_pct       = st.slider("Debt ratio (%)", 0, 90, 70, 5) / 100
    interest_pct   = st.slider("Interest rate (%)", 2.0, 10.0, 5.5, 0.5) / 100
    loan_years     = st.slider("Loan term (years)", 5, 20, 15, 1)
    tax_pct        = st.slider("Corporate tax rate (%)", 0, 40, 25, 1) / 100
    project_life   = 25


# ── GATE TABS ──────────────────────────────────────────────
tabs = st.tabs([
    "**Gate 1** Site",
    "**Gate 2** Aviation Safety",
    "**Gate 3** Grid",
    "**Gate 4** Commercial",
    "**Gate 5** GF vs BF",
    "**Gate 6** Sizing & Finance",
    "**Gate 7** Spatial Visuals",
    "**Export**",
])

# ── GATE 1 ─────────────────────────────────────────────────
with tabs[0]:
    st.markdown('<div class="section-label">Gate 1 — Site Viability</div>', unsafe_allow_html=True)
    st.markdown("Confirm that the site can physically accommodate solar before any further analysis.")

    c1, c2 = st.columns(2)
    with c1:
        total_site_m2      = st.number_input(
            "Total site area (m²)", 10_000.0, 100_000_000.0, 5_000_000.0, 50_000.0,
            format="%0.0f",
            help="1 ha = 10,000 m². A large hub is typically 5–20 million m²")
        st.caption(f"= {total_site_m2/10_000:,.1f} ha")
        terrain            = st.selectbox("Terrain type", [
            "Flat / optimal (multiplier: ×1.0)",
            "Gentle slope 0–5% (multiplier: ×1.05 CAPEX, ×1.10 area)",
            "Complex / steep >5% (multiplier: ×1.15 CAPEX, ×1.25 area)",
        ])
        land_ownership     = st.selectbox("Land ownership", [
            "Owned by airport – no acquisition cost",
            "Owned by airport – part leased to third party",
            "Must be acquired / leased",
        ])
    with c2:
        built_pct          = st.slider("Built / operational area (%)", 0, 80, 40, 5)
        green_protected_m2 = st.number_input(
            "Protected green / ecology area (m²)", 0.0, 50_000_000.0, 0.0, 50_000.0,
            format="%0.0f",
            help="Bird protection zones, flood plain, compensation habitat")
        committed_dev_m2   = st.number_input(
            "Area committed to other development (m²)", 0.0, 50_000_000.0, 0.0, 50_000.0,
            format="%0.0f",
            help="Land already allocated to terminal expansion, cargo or other projects")

    terrain_area_mult = 1.0
    terrain_capex_mult = 1.0
    if "Gentle" in terrain:
        terrain_area_mult = 1.10; terrain_capex_mult = 1.05
    elif "steep" in terrain:
        terrain_area_mult = 1.25; terrain_capex_mult = 1.15

    built_m2 = total_site_m2 * built_pct / 100
    gate1_available_m2 = total_site_m2 - built_m2 - green_protected_m2 - committed_dev_m2
    gate1_available_m2 = max(0.0, gate1_available_m2)

    # Acquisition cost
    land_cost_total = 0.0
    if "acquired" in land_ownership:
        land_cost_per_m2 = st.number_input(
                "Land acquisition cost (€ / m²)", 0.0, 500.0, 5.0, 0.5,
                help="€5/m² = €50,000/ha")
        land_cost_total = gate1_available_m2 * land_cost_per_m2

    # Flags
    g1_flags = []
    # Indicative capacity range across the technology options (7,000–30,000 m²/MW)
    mw_low  = gate1_available_m2 / 30_000
    mw_high = gate1_available_m2 /  7_000
    if gate1_available_m2 < 200_000:                      # under 20 ha
        g1_flags.append(("fail",
            f"{gate1_available_m2:,.0f} m² ({gate1_available_m2/10_000:,.1f} ha) available — "
            f"below the 200,000 m² threshold for a meaningful installation"))
    elif gate1_available_m2 < 500_000:                    # under 50 ha
        g1_flags.append(("warn",
            f"{gate1_available_m2:,.0f} m² ({gate1_available_m2/10_000:,.1f} ha) available — "
            f"supports roughly {mw_low:,.0f}–{mw_high:,.0f} MW depending on technology"))
    else:
        g1_flags.append(("pass",
            f"{gate1_available_m2:,.0f} m² ({gate1_available_m2/10_000:,.1f} ha) available — "
            f"sufficient for preliminary sizing"))
    if "acquired" in land_ownership and land_cost_total > 0:
        g1_flags.append(("warn",
            f"Land acquisition adds €{land_cost_total:,.0f} to CAPEX — verify against project economics"))
    if green_protected_m2 > 0:
        g1_flags.append(("info",
            f"{green_protected_m2:,.0f} m² excluded as protected — confirm the boundary with an ecology survey"))

    st.markdown('<hr class="m2p-divider">', unsafe_allow_html=True)
    st.markdown("**Gate 1 assessment**")
    for status, msg in g1_flags:
        badge = {"pass": "badge-pass", "warn": "badge-warn", "fail": "badge-fail", "info": "badge-info"}[status]
        label = status.upper()
        st.markdown(f'<span class="badge {badge}">{label}</span> {msg}', unsafe_allow_html=True)

    st.info(f"**Available area carried forward to Gate 2:** {gate1_available_m2:,.0f} m²  ({gate1_available_m2/10_000:,.1f} ha)")

# ── GATE 2 ─────────────────────────────────────────────────
with tabs[1]:
    st.markdown('<div class="section-label">Gate 2 — Aviation Safety Constraints</div>', unsafe_allow_html=True)
    st.markdown("Aviation safety constraints typically reduce usable area. Complete each check before committing to a layout.")

    c1, c2 = st.columns(2)
    with c1:
        glare_study_done  = st.checkbox("Formal glare study completed")
        glare_zones_m2    = st.number_input(
            "Area eliminated by glare constraints (m²)", 0.0, float(gate1_available_m2),
            0.0, 10_000.0, format="%0.0f",
            disabled=not glare_study_done,
            help="From the formal glare study report")
        ols_study_done    = st.checkbox("Obstacle limitation surface (OLS) mapping completed")
        ols_eliminated_m2 = st.number_input(
            "Area eliminated by OLS height restrictions (m²)", 0.0, float(gate1_available_m2),
            0.0, 10_000.0, format="%0.0f",
            disabled=not ols_study_done,
            help="Obstacle limitation surface height caps")
    with c2:
        rpz_eliminated_m2  = st.number_input(
            "Area within runway protection zones (m²)", 0.0, float(gate1_available_m2),
            0.0, 10_000.0, format="%0.0f",
            help="No solar permitted within an RPZ regardless of any other clearance")
        naa_consulted      = st.checkbox("Pre-application consultation with national aviation authority completed")
        antenna_ok         = st.checkbox("Confirmed no interference with navigation/communication equipment")

    g2_eliminated_m2 = glare_zones_m2 + ols_eliminated_m2 + rpz_eliminated_m2
    gate2_cleared_m2 = max(0.0, gate1_available_m2 - g2_eliminated_m2)

    # Flags
    g2_flags = []
    if not glare_study_done:
        g2_flags.append(("warn", "Glare study not yet completed — area available may be overstated"))
    if not ols_study_done:
        g2_flags.append(("warn", "OLS mapping not completed — height restrictions unknown"))
    if rpz_eliminated_m2 > 0:
        g2_flags.append(("fail", f"{rpz_eliminated_m2:,.0f} m² within a runway protection zone — this area is unavailable"))
    if not naa_consulted:
        g2_flags.append(("warn", "National aviation authority consultation required before design is committed"))
    if not antenna_ok:
        g2_flags.append(("fail", "Navigation/communication interference unresolved — do not proceed without clearance"))
    if gate2_cleared_m2 > 0 and len([f for f in g2_flags if f[0] == "fail"]) == 0:
        g2_flags.append(("pass", f"{gate2_cleared_m2:,.0f} m² ({gate2_cleared_m2/10_000:,.1f} ha) cleared for preliminary layout — subject to completing outstanding studies"))

    reduction_pct = (1 - gate2_cleared_m2 / gate1_available_m2) * 100 if gate1_available_m2 > 0 else 0

    st.markdown('<hr class="m2p-divider">', unsafe_allow_html=True)
    st.markdown("**Gate 2 assessment**")
    for status, msg in g2_flags:
        badge = {"pass": "badge-pass", "warn": "badge-warn", "fail": "badge-fail", "info": "badge-info"}[status]
        st.markdown(f'<span class="badge {badge}">{status.upper()}</span> {msg}', unsafe_allow_html=True)

    col_a, col_b = st.columns(2)
    col_a.metric("Aviation-cleared area", f"{gate2_cleared_m2:,.0f} m²",
                 help=f"= {gate2_cleared_m2/10_000:,.2f} ha")
    col_b.metric("Reduction from Gate 1", f"{reduction_pct:.0f}%",
                 delta=f"−{g2_eliminated_m2:,.0f} m²", delta_color="inverse")

# ── GATE 3 ─────────────────────────────────────────────────
with tabs[2]:
    st.markdown('<div class="section-label">Gate 3 — Grid Connection</div>', unsafe_allow_html=True)
    st.markdown("Solar is only useful if the grid can absorb it. Resolve grid capacity before sizing the system.")

    c1, c2 = st.columns(2)
    with c1:
        annual_demand_gwh   = st.number_input("Airport annual electricity demand (GWh/yr)", 0.1, 2000.0, 150.0, 10.0)
        grid_headroom_mw    = st.number_input("Available grid connection headroom (MW)", 0.0, 2000.0, 50.0, 5.0,
                                               help="From distribution network operator (DNO) pre-application enquiry")
        grid_study_done     = st.checkbox("DNO pre-application enquiry completed")
        export_allowed      = st.checkbox("Grid export permitted (excess generation can be sold)")
    with c2:
        self_consume_pct    = st.slider("Estimated self-consumption (%)", 10, 100, 70, 5,
                                         help="The share of generation consumed on site. Higher = better economics")
        add_bess            = st.checkbox("Include battery storage (BESS) to shift generation to peak demand?")
        bess_mwh            = 0.0
        bess_capex_mwh      = 250000.0
        bess_revenue_mwh_yr = 0.0
        if add_bess:
            bess_mwh            = st.number_input("Battery capacity (MWh)", 10.0, 500.0, 50.0, 10.0)
            bess_capex_mwh      = st.number_input("BESS CAPEX (€ / MWh)", 150000, 500000, 250000, 10000)
            bess_revenue_mwh_yr = st.number_input("BESS annual value (€ / MWh installed)", 0, 50000, 8000, 1000,
                                                   help="From arbitrage, capacity payments and ancillary services. Conservative = €5k–€10k/MWh/yr")

    g3_flags = []
    if not grid_study_done:
        g3_flags.append(("warn", "DNO pre-application enquiry not completed — grid headroom is unconfirmed"))
    if grid_headroom_mw < 10:
        g3_flags.append(("fail", "Grid headroom under 10 MW — grid upgrade required before proceeding"))
    elif grid_headroom_mw < 30:
        g3_flags.append(("warn", f"Grid headroom {grid_headroom_mw:.0f} MW — limits maximum capacity; consider BESS to reduce peak export"))
    else:
        g3_flags.append(("pass", f"Grid headroom {grid_headroom_mw:.0f} MW — sufficient for preliminary sizing"))
    if not export_allowed and self_consume_pct < 100:
        g3_flags.append(("warn", "Export not permitted — size system to self-consumption only, or generation will be curtailed"))

    st.markdown('<hr class="m2p-divider">', unsafe_allow_html=True)
    for status, msg in g3_flags:
        badge = {"pass": "badge-pass", "warn": "badge-warn", "fail": "badge-fail", "info": "badge-info"}[status]
        st.markdown(f'<span class="badge {badge}">{status.upper()}</span> {msg}', unsafe_allow_html=True)

# ── GATE 4 ─────────────────────────────────────────────────
with tabs[3]:
    st.markdown('<div class="section-label">Gate 4 — Commercial Structure</div>', unsafe_allow_html=True)
    st.markdown("The commercial model determines who builds it, who owns it, and who carries the risk.")

    commercial_model = st.selectbox("Commercial model", [
        "Direct ownership (airport owns and operates)",
        "Power Purchase Agreement – airport buys output from a third-party developer",
        "Special Purpose Vehicle / concession (developer owns, airport receives carbon benefit)",
        "Rooftop lease (developer pays rent, airport receives rental income)",
    ])

    c1, c2 = st.columns(2)
    with c1:
        airport_ownership  = st.selectbox("Airport ownership", [
            "Public (government / municipal)",
            "Private",
            "Mixed (public majority)",
        ])
        can_fund_capex     = st.checkbox("Airport can fund full CAPEX from own balance sheet")
        ppa_market_exists  = st.checkbox("A PPA market exists in this jurisdiction")
    with c2:
        permit_required    = st.checkbox("Generation / renewable energy licence required")
        permit_obtained    = st.checkbox("Licence obtained", disabled=not permit_required)
        aca_target         = st.checkbox("Airport has an ACA or equivalent carbon target")
        aca_level          = st.selectbox("ACA level (if applicable)", ["Not applicable", "Level 1", "Level 2",
                                                                         "Level 3", "Level 3+", "Level 4", "Level 4+", "Level 5"],
                                           disabled=not aca_target)

    model_notes = {
        "Direct ownership": "Airport carries full CAPEX. Best when balance sheet is strong and tax position allows depreciation benefits.",
        "Power Purchase": "Zero CAPEX for airport. Developer takes construction and operational risk. Airport pays for electricity at agreed rate.",
        "Special Purpose": "SPV carries CAPEX. Airport contributes land. Used when airport cannot fund capital — see KLIA / Cenergi precedent.",
        "Rooftop lease": "Simplest structure. Developer pays rent for roof space. Airport receives income but no direct carbon benefit in the energy balance.",
    }
    matched_note = next((v for k, v in model_notes.items() if k in commercial_model), "")

    g4_flags = []
    if "Direct" in commercial_model and not can_fund_capex:
        g4_flags.append(("warn", "Direct ownership selected but airport cannot fund CAPEX — consider SPV or PPA route"))
    if "PPA" in commercial_model and not ppa_market_exists:
        g4_flags.append(("fail", "PPA selected but no PPA market in this jurisdiction — developer cannot be remunerated"))
    if permit_required and not permit_obtained:
        g4_flags.append(("warn", "Generation licence required but not yet obtained — add 6–12 months to programme"))
    if not g4_flags:
        g4_flags.append(("pass", f"Commercial model selected: {commercial_model.split('(')[0].strip()}"))

    if matched_note:
        st.info(f"**Model note:** {matched_note}")

    st.markdown('<hr class="m2p-divider">', unsafe_allow_html=True)
    for status, msg in g4_flags:
        badge = {"pass": "badge-pass", "warn": "badge-warn", "fail": "badge-fail", "info": "badge-info"}[status]
        st.markdown(f'<span class="badge {badge}">{status.upper()}</span> {msg}', unsafe_allow_html=True)

# ── GATE 5 ─────────────────────────────────────────────────
with tabs[4]:
    st.markdown('<div class="section-label">Gate 5 — Greenfield vs Brownfield Checklist</div>', unsafe_allow_html=True)
    st.markdown(f"Additional checks for **{airport_type}** developments.")

    if airport_type == "Greenfield":
        st.markdown("##### Greenfield checks")
        checks_gf = {
            "Solar integrated into master plan (not added after layout decisions)":
                "Retrofitting solar after layout is fixed can reduce yield 15–20% due to suboptimal orientation.",
            "Terminal roof pitch and orientation optimised for solar at this stage":
                "This is an architectural decision. Once fixed it cannot be changed without redesign.",
            "Ground-mount land safeguarded before parcels are allocated to other uses":
                "Once allocated, land is gone. Safeguarding costs nothing at masterplan stage.",
            "Grid connection designed to include solar generation from day one":
                "Retrofit grid upgrade can cost 2–3× a designed-in connection.",
            "Approach paths and obstacle surfaces designed to leave solar zones clear":
                "Cheaper to plan around them than conduct glare studies on a fixed layout.",
        }
        results_gf = {}
        for check, detail in checks_gf.items():
            results_gf[check] = st.checkbox(check, help=detail)
        incomplete_gf = [c for c, v in results_gf.items() if not v]
        if incomplete_gf:
            st.warning(f"**{len(incomplete_gf)} check(s) outstanding** — resolve before concept design is frozen.")
        else:
            st.success("All greenfield checks passed.")

    else:  # Brownfield
        st.markdown("##### Brownfield checks")
        checks_bf = {
            "Structural survey of existing roofs confirms load-bearing capacity for panels":
                "Standard panels add 15–20 kg/m². Many older terminal roofs cannot take this.",
            "Roof remaining life matches or exceeds 25-year panel life":
                "Replacing the roof mid-life requires full removal and reinstallation of panels.",
            "Shadow study completed for buildings, stands and parked aircraft":
                "Aircraft shadows can reduce rooftop yield by 20–30% in some configurations.",
            "Glare study covers all existing approach paths and taxiways (including any new ones)":
                "New taxiways or stand reconfigurations since original design may have changed the picture.",
            "Confirmed whether a land bank exists that is not committed to other development":
                "The best brownfield solar sites are often surplus operational land, not rooftops.",
            "Grid connection upgrade cost and lead time obtained from DNO":
                "Lead time can exceed 24 months and can be the longest item on the programme.",
        }
        results_bf = {}
        for check, detail in checks_bf.items():
            results_bf[check] = st.checkbox(check, help=detail)
        incomplete_bf = [c for c, v in results_bf.items() if not v]
        if incomplete_bf:
            st.warning(f"**{len(incomplete_bf)} check(s) outstanding** — resolve before committing to a layout.")
        else:
            st.success("All brownfield checks passed.")

# ── GATE 6 ─────────────────────────────────────────────────
with tabs[5]:
    st.markdown('<div class="section-label">Gate 6 — Indicative Sizing & Financial Model</div>', unsafe_allow_html=True)
    st.markdown("Run this only once Gates 1–5 are substantially clear. All figures are indicative only — not for final investment decisions.")

    c1, c2, c3 = st.columns(3)
    with c1:
        install_type       = st.selectbox("Installation type", [
            "Ground-mount – fixed tilt",
            "Ground-mount – single-axis tracking",
            "Rooftop – flat",
            "Rooftop – pitched / optimised",
            "Mixed (ground-mount + rooftop)",
        ])
        # Area per MW in m². A ~600 W module occupies ~2.4 m², so roughly 250 W/m²
        # of panel area — about 4,000 m² of pure module per MW before any walkways,
        # setbacks, tilt frames or inter-row spacing are added.
        m2_per_mw_lookup = {
            "Ground-mount – fixed tilt":            20_000,  # incl. inter-row spacing
            "Ground-mount – single-axis tracking":  30_000,  # wider rows, avoids shading
            "Rooftop – flat":                       12_000,  # tilt frames + row spacing
            "Rooftop – pitched / optimised":         7_000,  # flush-mounted, follows pitch
            "Mixed (ground-mount + rooftop)":       18_000,
        }
        m2_per_mw          = m2_per_mw_lookup[install_type] * terrain_area_mult
        usable_fraction    = st.slider("Usable fraction after setbacks / shading (%)", 40, 90, 70, 5) / 100
    with c2:
        effective_m2       = gate2_cleared_m2 * usable_fraction
        indicative_mw      = effective_m2 / m2_per_mw
        capacity_override  = st.checkbox("Override with a target capacity (client-specified)")
        if capacity_override:
            indicative_mw  = st.number_input("Target capacity (MW)", 1.0, 1000.0,
                                              float(max(1.0, round(indicative_mw, 0))), 1.0)
    with c3:
        system_eff         = st.slider("System efficiency (%)", 70, 90, 80, 1) / 100
        degradation_yr     = st.slider("Annual degradation (%)", 0.3, 1.0, 0.5, 0.1) / 100

    # ── CALCULATION ENGINE ──────────────────────────────────
    capex_solar         = indicative_mw * capex_mw * terrain_capex_mult
    capex_bess          = bess_mwh * bess_capex_mwh if add_bess else 0.0
    capex_land          = land_cost_total if "acquired" in land_ownership else 0.0
    total_capex         = capex_solar + capex_bess + capex_land
    depreciable_base    = capex_solar + capex_bess       # land NOT depreciated
    equity_invest       = total_capex * (1 - debt_pct)
    debt_principal      = total_capex * debt_pct

    # ── THERMAL & BIFACIAL MODEL ────────────────────────────
    # Monthly ambient profile: sinusoid between summer peak and winter minimum.
    # Month 0 = January (northern hemisphere winter).
    MONTHS = ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]
    amb_min   = amb_temp_peak - amb_temp_swing
    amb_mean  = (amb_temp_peak + amb_min) / 2
    amb_amp   = (amb_temp_peak - amb_min) / 2
    # Peak in July (index 6): cos curve shifted so month 6 is maximum
    monthly_amb = [amb_mean + amb_amp * math.cos(2 * math.pi * (m - 6) / 12) for m in range(12)]

    # Monthly irradiance profile: scaled from annual peak sun hours, same seasonal shape
    # but flatter (solar resource varies less than temperature in low latitudes).
    irr_amp_frac = 0.25          # ±25% seasonal variation in daily PSH
    monthly_psh  = [peak_sun * (1 + irr_amp_frac * math.cos(2 * math.pi * (m - 6) / 12))
                    for m in range(12)]

    # Rear-side view factor by installation type — clearance beneath the array
    # is what makes bifacial work. Flush-mounted roofs get almost nothing.
    view_factor_lookup = {
        "Ground-mount – fixed tilt":            0.35,
        "Ground-mount – single-axis tracking":  0.40,
        "Rooftop – flat":                       0.12,
        "Rooftop – pitched / optimised":        0.08,
        "Mixed (ground-mount + rooftop)":       0.25,
    }
    view_factor = view_factor_lookup.get(install_type, 0.20)

    # NOCT cell-temperature model (IEC 61215 reference conditions: 800 W/m², 20°C, 1 m/s)
    # T_cell = T_ambient + (NOCT − 20) / 800 × G
    G_PEAK = 800.0               # W/m² reference irradiance for the NOCT relation

    monthly_rows = []
    for m in range(12):
        t_amb    = monthly_amb[m]
        psh_m    = monthly_psh[m]
        t_cell   = t_amb + (noct - 20.0) / 800.0 * G_PEAK
        # Derate factor: 1 + coeff × (T_cell − 25). coeff is negative.
        derate   = 1.0 + temp_coeff * (t_cell - 25.0)
        derate   = max(derate, 0.0)
        # Bifacial rear-side gain — additive irradiance, independent of thermal loss
        bi_gain  = (albedo * bifaciality * view_factor) if bifacial else 0.0

        days     = 30.42         # mean days per month
        kwp      = indicative_mw * 1000
        std_kwh  = kwp * psh_m * system_eff * days
        temp_kwh = std_kwh * derate
        final_kwh= temp_kwh * (1.0 + bi_gain)

        monthly_rows.append({
            "Month": MONTHS[m],
            "Ambient (°C)": round(t_amb, 1),
            "Cell temp (°C)": round(t_cell, 1),
            "Derate factor": round(derate, 4),
            "Standard yield (MWh)": std_kwh / 1000,
            "Temp-adjusted (MWh)": temp_kwh / 1000,
            "Final yield (MWh)": final_kwh / 1000,
        })

    df_month = pd.DataFrame(monthly_rows)

    std_annual_kwh   = df_month["Standard yield (MWh)"].sum() * 1000
    temp_annual_kwh  = df_month["Temp-adjusted (MWh)"].sum() * 1000
    base_kwh         = df_month["Final yield (MWh)"].sum() * 1000   # feeds the financial model

    heat_loss_pct    = (1 - temp_annual_kwh / std_annual_kwh) * 100 if std_annual_kwh else 0.0
    bifacial_gain_pct= (base_kwh / temp_annual_kwh - 1) * 100 if temp_annual_kwh else 0.0
    net_vs_std_pct   = (base_kwh / std_annual_kwh - 1) * 100 if std_annual_kwh else 0.0

    peak_cell_temp   = df_month["Cell temp (°C)"].max()
    worst_month      = df_month.loc[df_month["Derate factor"].idxmin(), "Month"]
    worst_derate_pct = (1 - df_month["Derate factor"].min()) * 100

    annual_co2_t     = base_kwh * 0.233 / 1000       # EU average emission factor kg/kWh

    if debt_pct > 0 and interest_pct > 0 and loan_years > 0:
        annual_debt_svc = (debt_principal * interest_pct) / (1 - (1 + interest_pct) ** -loan_years)
    else:
        annual_debt_svc = 0.0
    annual_dep          = depreciable_base / project_life
    outstanding_debt    = debt_principal
    cash_flows          = [-equity_invest]
    records             = []

    for yr in range(1, project_life + 1):
        yr_kwh      = base_kwh * ((1 - degradation_yr) ** (yr - 1))
        rev_solar   = yr_kwh * ppa_price * ((1 + ppa_escalation) ** (yr - 1)) * (self_consume_pct / 100)
        rev_bess    = (bess_mwh * bess_revenue_mwh_yr) if add_bess else 0.0
        revenue     = rev_solar + rev_bess
        opex        = (capex_solar * opex_pct) * ((1 + opex_inflation) ** (yr - 1))
        lifecycle   = (indicative_mw * 60000) if yr == 15 else 0.0   # inverter replacement, capitalised separately

        if yr <= loan_years:
            interest = outstanding_debt * interest_pct
            principal_pay = annual_debt_svc - interest
            debt_svc  = annual_debt_svc
            outstanding_debt = max(0, outstanding_debt - principal_pay)
        else:
            interest = 0.0; debt_svc = 0.0

        taxable     = revenue - opex - annual_dep - interest
        if yr == 15:
            taxable -= lifecycle * 0.03   # partial expense allowance
        tax         = max(0.0, taxable * tax_pct)
        net_cf      = revenue - opex - lifecycle - debt_svc - tax

        cash_flows.append(net_cf)
        records.append({"Year": yr, "Revenue": revenue, "OPEX": -opex,
                         "Debt Service": -debt_svc, "Tax": -tax,
                         "Net Cash Flow": net_cf})

    df = pd.DataFrame(records)
    df["Cumulative"] = df["Net Cash Flow"].cumsum() - equity_invest

    # IRR
    irr_val = npf.irr(cash_flows)
    irr_str = f"{irr_val:.1%}" if (irr_val is not None and not (irr_val != irr_val)) else "Not achievable at these inputs"

    pb_series = df[df["Cumulative"] > 0]["Year"]
    pb_str = f"Year {int(pb_series.min())}" if not pb_series.empty else "Beyond 25 years"

    # ── KPI TILES ───────────────────────────────────────────
    st.markdown('<hr class="m2p-divider">', unsafe_allow_html=True)
    st.markdown(f"### {project_name} — Indicative results")

    st.markdown(f"""
    <div class="kpi-row">
      <div class="kpi-tile blue">
        <div class="val">{indicative_mw:.1f} MW</div>
        <div class="lbl">Indicative capacity</div>
      </div>
      <div class="kpi-tile">
        <div class="val">{effective_m2:,.0f} m²</div>
        <div class="lbl">Effective PV area</div>
      </div>
      <div class="kpi-tile teal">
        <div class="val">{base_kwh/1e6:.1f} GWh</div>
        <div class="lbl">Year 1 generation</div>
      </div>
      <div class="kpi-tile green">
        <div class="val">{annual_co2_t:,.0f} t</div>
        <div class="lbl">CO₂ avoided / yr</div>
      </div>
      <div class="kpi-tile amber">
        <div class="val">€{total_capex/1e6:.1f}M</div>
        <div class="lbl">Total CAPEX</div>
      </div>
      <div class="kpi-tile">
        <div class="val">{irr_str}</div>
        <div class="lbl">After-tax levered IRR</div>
      </div>
      <div class="kpi-tile blue">
        <div class="val">{pb_str}</div>
        <div class="lbl">Equity payback</div>
      </div>
    </div>
    """, unsafe_allow_html=True)


    # ── CALCULATION METHODOLOGY — SIZING & FINANCE ──────────
    with st.expander("View Calculation Methodology — sizing and financials"):

        st.markdown("##### Step 1 — From site area to installed capacity")
        st.latex(r"A_{effective} = A_{cleared} \times f_{usable}")
        st.latex(r"P_{MW} = \frac{A_{effective}\,[m^2]}{r_{m^2/MW} \times m_{terrain}}")
        st.markdown(f"""
| Term | Value | Source |
|---|---|---|
| `A_cleared` — aviation-cleared area | **{gate2_cleared_m2:,.0f} m²** | Gate 2 output |
| `f_usable` — usable fraction | **{usable_fraction:.0%}** | User input, Gate 6 |
| `A_effective` — effective area | **{effective_m2:,.0f} m²** | {gate2_cleared_m2:,.0f} × {usable_fraction:.2f} |
| `r_m²/MW` — area per MW | **{m2_per_mw_lookup[install_type]:,.0f} m²/MW** | {install_type} |
| `m_terrain` — terrain multiplier | **×{terrain_area_mult:.2f}** | {terrain.split('(')[0].strip()} |
| **`P_MW` — indicative capacity** | **{indicative_mw:,.2f} MW** | {effective_m2:,.0f} ÷ ({m2_per_mw_lookup[install_type]:,.0f} × {terrain_area_mult:.2f}) |
""")
        if capacity_override:
            st.warning(f"Capacity override is active. The area-derived figure has been replaced "
                       f"by the client-specified target of {indicative_mw:,.1f} MW. The area check "
                       f"above no longer constrains the result.")

        st.markdown("##### Step 2 — From capacity to annual generation")
        st.latex(r"E_{month} = P_{kWp} \times H_{month} \times \eta_{sys} \times d \times D_{temp} \times (1 + g_{bifacial})")
        st.markdown(f"""
| Term | Value | Note |
|---|---|---|
| `P_kWp` — installed capacity | **{indicative_mw * 1000:,.0f} kWp** | {indicative_mw:,.2f} MW x 1,000 |
| `H_month` — peak sun hours/day | **{peak_sun:.2f} h** annual mean | Varies +/-{irr_amp_frac:.0%} seasonally |
| `eta_sys` — system efficiency | **{system_eff:.0%}** | Inverter, wiring, mismatch losses |
| `d` — days per month | **30.42** | 365 / 12 |
| `D_temp` — thermal derate | **{df_month['Derate factor'].mean():.4f}** mean | See thermal methodology below |
| `g_bifacial` — rear-side gain | **{(albedo * bifaciality * view_factor) if bifacial else 0:.4f}** | {"Bifacial enabled" if bifacial else "Not enabled"} |
| **Year 1 generation** | **{base_kwh/1e6:,.2f} GWh** | Sum of 12 monthly values |
""")

        st.markdown("##### Step 3 — Capital cost build-up")
        st.latex(r"CAPEX_{total} = (P_{MW} \times c_{MW} \times m_{terrain}) + (E_{BESS} \times c_{BESS}) + C_{land}")
        st.markdown(f"""
| Component | Calculation | Value |
|---|---|---|
| Solar | {indicative_mw:,.2f} MW x EUR{capex_mw:,.0f}/MW x {terrain_capex_mult:.2f} | **EUR{capex_solar:,.0f}** |
| Battery storage | {bess_mwh:,.1f} MWh x EUR{bess_capex_mwh:,.0f}/MWh | **EUR{capex_bess:,.0f}** |
| Land acquisition | {"Airport-owned — no cost" if capex_land == 0 else f"{gate1_available_m2:,.0f} m² acquired"} | **EUR{capex_land:,.0f}** |
| **Total CAPEX** | | **EUR{total_capex:,.0f}** |
| Debt ({debt_pct:.0%}) | EUR{total_capex:,.0f} x {debt_pct:.2f} | EUR{debt_principal:,.0f} |
| **Equity required** | EUR{total_capex:,.0f} x {1-debt_pct:.2f} | **EUR{equity_invest:,.0f}** |
""")

        st.markdown("##### Step 4 — Annual cash flow, per year of the 25-year model")
        st.latex(r"CF_{net} = R - OPEX - L - DS - T")
        st.markdown(f"""
**Revenue** `R` = generation x PPA price x escalation x self-consumption share
Year 1: {base_kwh/1e6:,.2f} GWh x EUR{ppa_price:.3f}/kWh x {self_consume_pct}% = **EUR{base_kwh * ppa_price * (self_consume_pct/100):,.0f}**
{"Plus BESS value: " + f"{bess_mwh:,.0f} MWh x EUR{bess_revenue_mwh_yr:,.0f}/MWh = EUR{bess_mwh * bess_revenue_mwh_yr:,.0f}" if add_bess else ""}

**Operating cost** `OPEX` = solar CAPEX x {opex_pct:.1%}, inflating at {opex_inflation:.1%}/yr
Year 1: EUR{capex_solar:,.0f} x {opex_pct:.3f} = **EUR{capex_solar * opex_pct:,.0f}**

**Lifecycle cost** `L` = inverter replacement in year 15 only, at EUR60,000/MW
= {indicative_mw:,.2f} MW x EUR60,000 = **EUR{indicative_mw * 60000:,.0f}** (year 15)

**Debt service** `DS` = level annuity over {loan_years} years at {interest_pct:.2%}
= **EUR{annual_debt_svc:,.0f}/yr** for years 1-{loan_years}

**Tax** `T` = max(0, (R - OPEX - depreciation - interest) x {tax_pct:.0%})
Straight-line depreciation = EUR{depreciable_base:,.0f} / 25 = **EUR{annual_dep:,.0f}/yr**
Note: land is excluded from the depreciable base.

**Degradation:** output falls {degradation_yr:.2%} per year, compounding.
""")

        st.markdown("##### Step 5 — Return metrics")
        st.markdown(f"""
- **IRR** is the discount rate at which the net present value of the equity cash flow series equals zero. The series begins with **-EUR{equity_invest:,.0f}** at year 0, followed by 25 annual net cash flows. Result: **{irr_str}**
- **Payback** is the first year in which cumulative cash flow turns positive. Result: **{pb_str}**
- **CO2 avoided** = {base_kwh:,.0f} kWh x 0.233 kg/kWh / 1,000 = **{annual_co2_t:,.0f} t/yr**
""")
        st.caption("All figures are pre-feasibility estimates. See the assumptions and limitations "
                   "section below before using any output externally.")

    # ── THERMAL PERFORMANCE PANEL ───────────────────────────
    st.markdown('<hr class="m2p-divider">', unsafe_allow_html=True)
    st.markdown('<div class="section-label">Hot-climate thermal performance</div>', unsafe_allow_html=True)

    bifacial_tile = (
        f'<div class="kpi-tile green"><div class="val">+{bifacial_gain_pct:.1f}%</div>'
        f'<div class="lbl">Yield recovered via bifacial</div></div>'
        if bifacial else
        f'<div class="kpi-tile" style="background:#8A8A8A"><div class="val">—</div>'
        f'<div class="lbl">Bifacial not selected</div></div>'
    )

    st.markdown(f"""
    <div class="kpi-row">
      <div class="kpi-tile amber">
        <div class="val">−{heat_loss_pct:.1f}%</div>
        <div class="lbl">Efficiency lost to heat</div>
      </div>
      {bifacial_tile}
      <div class="kpi-tile teal">
        <div class="val">{peak_cell_temp:.0f}°C</div>
        <div class="lbl">Peak cell temperature</div>
      </div>
      <div class="kpi-tile blue">
        <div class="val">{net_vs_std_pct:+.1f}%</div>
        <div class="lbl">Net vs standard yield</div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # Contextual warnings
    if heat_loss_pct > 12:
        st.warning(f"**Significant thermal derating.** At {amb_temp_peak}°C ambient the array loses "
                   f"{heat_loss_pct:.1f}% of standard yield across the year, peaking in {worst_month} "
                   f"at {worst_derate_pct:.1f}%. Specify modules with a temperature coefficient better "
                   f"than −0.35%/°C and ensure adequate rear ventilation.")
    if bifacial and view_factor < 0.15:
        st.warning(f"**Bifacial gain is limited on this installation type.** {install_type} gives a "
                   f"rear-side view factor of {view_factor:.2f}, so the bifacial premium may not be "
                   f"justified. Bifacial performs best on elevated ground-mount over reflective ground.")
    if bifacial and albedo > 0.5 and "Rooftop" in install_type:
        st.info("High albedo assumed on a rooftop — verify the roof membrane is genuinely light-coloured "
                "and will stay that way. Soiling reduces albedo substantially in dusty climates.")

    # ── CHARTS ──────────────────────────────────────────────
    col_chart1, col_chart2 = st.columns(2)

    # Monthly yield: standard vs temperature-adjusted vs final
    fig_m = go.Figure()
    fig_m.add_bar(x=df_month["Month"], y=df_month["Standard yield (MWh)"],
                  name="Standard yield (no thermal derating)",
                  marker_color="#B8B8B8", opacity=0.85)
    fig_m.add_bar(x=df_month["Month"], y=df_month["Temp-adjusted (MWh)"],
                  name="Temperature-adjusted yield",
                  marker_color=M2P_AMBER, opacity=0.9)
    if bifacial:
        fig_m.add_bar(x=df_month["Month"], y=df_month["Final yield (MWh)"],
                      name="With bifacial rear-side gain",
                      marker_color=M2P_GREEN, opacity=0.9)
    fig_m.add_scatter(x=df_month["Month"], y=df_month["Cell temp (°C)"],
                      name="Cell temperature (°C)", yaxis="y2",
                      mode="lines+markers", line=dict(color=M2P_NAVY, width=2, dash="dot"))
    fig_m.update_layout(
        title="Monthly yield — standard vs temperature-adjusted",
        barmode="group",
        xaxis_title="Month",
        yaxis=dict(title="Yield (MWh)"),
        yaxis2=dict(title="Cell temperature (°C)", overlaying="y", side="right",
                    showgrid=False, range=[0, max(100, peak_cell_temp * 1.3)]),
        legend=dict(orientation="h", y=-0.22),
        plot_bgcolor="white", paper_bgcolor="white",
        font=dict(family="Century Gothic, Trebuchet MS, sans-serif", size=11),
        margin=dict(l=50, r=60, t=50, b=60),
        height=430,
    )
    st.plotly_chart(fig_m, use_container_width=True)

    with st.expander("View monthly thermal calculation table"):
        show = df_month.copy()
        for c in ["Standard yield (MWh)", "Temp-adjusted (MWh)", "Final yield (MWh)"]:
            show[c] = show[c].apply(lambda v: f"{v:,.0f}")
        show["Loss to heat (%)"] = df_month["Derate factor"].apply(lambda d: f"{(1-d)*100:.1f}%")
        st.dataframe(show, use_container_width=True, hide_index=True)
        st.caption(
            f"Cell temperature from the NOCT model: T_cell = T_ambient + (NOCT − 20) / 800 × 800 W/m², "
            f"with NOCT = {noct}°C. Derating applies the {temp_coeff*100:.2f}%/°C coefficient to the "
            f"difference between cell temperature and the 25°C STC reference."
            + (f" Bifacial gain = albedo {albedo:.2f} × bifaciality {bifaciality:.2f} × "
               f"view factor {view_factor:.2f} = +{albedo*bifaciality*view_factor*100:.1f}%."
               if bifacial else "")
        )


    with st.expander("View Calculation Methodology — monthly thermal profile"):

        st.markdown("##### Step 1 — Monthly ambient temperature")
        st.latex(r"T_{amb}(m) = T_{mean} + A \times \cos\!\left(\frac{2\pi(m-6)}{12}\right)")
        st.markdown(f"""
A sinusoid with its maximum in July (month index 6).

| Term | Value |
|---|---|
| Peak ambient (summer) | **{amb_temp_peak} degC** |
| Annual swing | **{amb_temp_swing} degC** |
| Winter minimum | {amb_temp_peak} - {amb_temp_swing} = **{amb_temp_peak - amb_temp_swing} degC** |
| `T_mean` — annual mean | **{amb_mean:.1f} degC** |
| `A` — amplitude | **{amb_amp:.1f} degC** |
""")

        st.markdown("##### Step 2 — Cell temperature from ambient")
        st.latex(r"T_{cell} = T_{amb} + \frac{NOCT - 20}{800} \times G")
        st.markdown(f"""
The NOCT relation (IEC 61215 reference conditions: 800 W/m2 irradiance, 20 degC ambient, 1 m/s wind).

| Term | Value | Note |
|---|---|---|
| `NOCT` | **{noct} degC** | Module datasheet |
| `G` | **800 W/m2** | Reference irradiance |
| Temperature rise above ambient | ({noct} - 20) / 800 x 800 = **{noct - 20:.0f} degC** | Constant at reference irradiance |
| Peak cell temperature | {amb_temp_peak} + {noct - 20:.0f} = **{peak_cell_temp:.0f} degC** | July |
""")

        st.markdown("##### Step 3 — Thermal derating")
        st.latex(r"D_{temp} = 1 + \gamma \times (T_{cell} - 25)")
        st.markdown(f"""
| Term | Value |
|---|---|
| `gamma` — temperature coefficient | **{temp_coeff*100:.2f} %/degC** |
| STC reference temperature | **25 degC** |
| Worst month | **{worst_month}** — derate {df_month['Derate factor'].min():.4f}, a **{worst_derate_pct:.1f}%** loss |
| Best month | **{df_month.loc[df_month['Derate factor'].idxmax(), 'Month']}** — derate {df_month['Derate factor'].max():.4f} |
| **Annual weighted loss** | **{heat_loss_pct:.1f}%** |
""")

        st.markdown("##### Step 4 — Bifacial rear-side gain")
        if bifacial:
            st.latex(r"g_{bifacial} = \rho \times \beta \times VF")
            st.markdown(f"""
| Term | Value | Source |
|---|---|---|
| `rho` — surface albedo | **{albedo:.2f}** | User input |
| `beta` — bifaciality factor | **{bifaciality:.2f}** | Module datasheet |
| `VF` — rear-side view factor | **{view_factor:.2f}** | Set by installation type: {install_type} |
| **Gain** | **+{albedo * bifaciality * view_factor * 100:.1f}%** | {albedo:.2f} x {bifaciality:.2f} x {view_factor:.2f} |
""")
            st.info("Bifacial gain and thermal derating are independent physical effects. "
                    "Bifacial adds rear-side irradiance; it does not recover heat losses. "
                    "The two are reported separately for that reason.")
        else:
            st.markdown("Bifacial modules are not selected, so no rear-side gain is applied.")

        st.markdown("##### Step 5 — Reconciliation")
        st.markdown(f"""
| Stage | Annual yield | Change |
|---|---|---|
| Standard (no derating) | **{std_annual_kwh/1e6:,.2f} GWh** | baseline |
| After thermal derating | **{temp_annual_kwh/1e6:,.2f} GWh** | **-{heat_loss_pct:.1f}%** |
| After bifacial gain | **{base_kwh/1e6:,.2f} GWh** | **+{bifacial_gain_pct:.1f}%** |
| **Net versus standard** | | **{net_vs_std_pct:+.1f}%** |
""")
        st.warning("Soiling is not included in these figures. In Gulf conditions dust reduces "
                   "output by a further 15-30% without a scheduled cleaning programme.")

    with col_chart1:
        fig1 = go.Figure()
        fig1.add_bar(x=df["Year"], y=df["Net Cash Flow"]/1e6,
                     name="Annual net cash flow", marker_color=M2P_BLUE, opacity=0.75)
        fig1.add_scatter(x=df["Year"], y=df["Cumulative"]/1e6,
                         mode="lines+markers", name="Cumulative return",
                         line=dict(color=M2P_AMBER, width=2.5))
        fig1.add_hline(y=0, line_dash="dot", line_color="grey", opacity=0.5)
        fig1.update_layout(
            title="Cash flows over project life",
            xaxis_title="Year", yaxis_title="€ million",
            legend=dict(orientation="h", y=-0.2),
            plot_bgcolor="white", paper_bgcolor="white",
            font=dict(family="Century Gothic, Trebuchet MS, sans-serif", size=11),
            margin=dict(l=40, r=20, t=50, b=40),
        )
        st.plotly_chart(fig1, use_container_width=True)

    with col_chart2:
        capex_breakdown = {
            "Solar panels & BOS": capex_solar / 1e6,
            "Battery storage (BESS)": capex_bess / 1e6,
            "Land acquisition": capex_land / 1e6,
        }
        capex_breakdown = {k: v for k, v in capex_breakdown.items() if v > 0}
        fig2 = go.Figure(go.Pie(
            labels=list(capex_breakdown.keys()),
            values=list(capex_breakdown.values()),
            marker_colors=[M2P_BLUE, M2P_TEAL, M2P_AMBER],
            textinfo="label+percent",
            hole=0.42,
        ))
        fig2.update_layout(
            title="CAPEX breakdown",
            font=dict(family="Century Gothic, Trebuchet MS, sans-serif", size=11),
            paper_bgcolor="white",
            margin=dict(l=20, r=20, t=50, b=20),
        )
        st.plotly_chart(fig2, use_container_width=True)

    # ── DETAILED TABLE ──────────────────────────────────────
    with st.expander("View year-by-year cash flow table"):
        display_df = df.copy()
        for col in ["Revenue", "OPEX", "Debt Service", "Tax", "Net Cash Flow", "Cumulative"]:
            display_df[col] = display_df[col].apply(lambda x: f"€{x:,.0f}")
        st.dataframe(display_df, use_container_width=True, hide_index=True)

    # ── BENCHMARKS ──────────────────────────────────────────
    st.markdown('<hr class="m2p-divider">', unsafe_allow_html=True)
    st.markdown('<div class="section-label">Calibration against M2P case studies</div>', unsafe_allow_html=True)
    bench_data = {
        "Airport":      ["Munich (MUC)",         "KLIA (KUL)",                      "SFO"],
        "Country":      ["Germany",              "Malaysia",                         "USA"],
        "Capacity":     ["50 MWp (target 2030)", "30 MW (due 2027)",                "~5 MW"],
        "Annual yield": ["~45 GWh",              "~46 GWh",                          "~8 GWh"],
        "CO₂ avoided":  ["~18,000 t/yr",         "~35,000 t/yr",                    "~1,800 t/yr"],
        "Structure":    ["Direct ownership",     "SPV with Cenergi SEA",             "Grid procurement"],
        "Notes":        ["14% complete; 1,600 ha site; 2/3 protected",
                         "Airport land bank; no balance sheet cost",
                         "1.6% of 290 GWh demand; procurement dominant"],
    }
    st.dataframe(pd.DataFrame(bench_data), use_container_width=True, hide_index=True)

    # ── ASSUMPTION FLAGS ────────────────────────────────────
    st.markdown('<hr class="m2p-divider">', unsafe_allow_html=True)
    st.markdown('<div class="section-label">Model assumptions and limitations</div>', unsafe_allow_html=True)
    assumptions = [
        ("info", "Cell temperature uses the NOCT model at 800 W/m² reference irradiance. Actual cell temperature varies with wind speed, mounting configuration and rear ventilation — a detailed thermal model requires site meteorological data."),
        ("info", "Monthly ambient temperature is modelled as a sinusoid between the peak and the peak minus the annual swing, with the maximum in July. Replace with local monthly mean temperatures for a site-specific assessment."),
        ("warn", "Bifacial gain and thermal derating are independent effects. Bifacial adds rear-side irradiance; it does not recover heat losses. The two are shown separately for that reason."),
        ("warn", "Soiling losses are NOT modelled. In Gulf conditions dust accumulation reduces output 15–30% without a regular cleaning programme. Add this separately when assessing a desert site."),
        ("info", "Battery storage revenue is modelled as a user-input annual value per MWh. A full BESS dispatch model requires a separate analysis."),
        ("info", "CO₂ avoided uses the EU average grid emission factor (0.233 kg/kWh). Use a local factor for non-EU sites."),
        ("info", "Land cost is zero where airport owns the site. Update the acquisition figure at Gate 1 if purchasing land."),
        ("warn", "IRR and payback are pre-feasibility estimates. A full financial model for investment decisions requires audited demand data, a formal grid study, and tax advice in the relevant jurisdiction."),
        ("warn", "Glare and OLS assessments must be completed before any layout is committed. Uncompleted studies are flagged at Gate 2."),
    ]
    for status, msg in assumptions:
        badge = {"warn": "badge-warn", "info": "badge-info"}[status]
        st.markdown(f'<span class="badge {badge}">{status.upper()}</span> {msg}', unsafe_allow_html=True)

# ── GATE 7 — SPATIAL VISUALS ───────────────────────────────
with tabs[6]:
    st.markdown('<div class="section-label">Gate 7 — Spatial Visualisation</div>',
                unsafe_allow_html=True)
    st.markdown("How the site reduces from total available land to the area that can actually "
                "carry panels, and what that footprint looks like at true geographic scale.")

    # ══ 7.1 WATERFALL ═══════════════════════════════════════
    st.markdown("#### Area reduction waterfall")

    layout_loss_m2 = max(0.0, gate2_cleared_m2 - effective_m2)

    wf_labels = [
        "Available area<br>(Gate 1)",
        "Glare zones",
        "OLS restrictions",
        "Runway protection",
        "Layout &amp; setbacks",
        "Effective PV area",
    ]
    wf_measure = ["absolute", "relative", "relative", "relative", "relative", "total"]
    wf_values = [
        gate1_available_m2,
        -glare_zones_m2,
        -ols_eliminated_m2,
        -rpz_eliminated_m2,
        -layout_loss_m2,
        0,
    ]

    fig_wf = go.Figure(go.Waterfall(
        orientation="v",
        measure=wf_measure,
        x=wf_labels,
        y=wf_values,
        text=[f"{gate1_available_m2:,.0f}",
              f"−{glare_zones_m2:,.0f}",
              f"−{ols_eliminated_m2:,.0f}",
              f"−{rpz_eliminated_m2:,.0f}",
              f"−{layout_loss_m2:,.0f}",
              f"{effective_m2:,.0f}"],
        textposition="outside",
        textfont=dict(family="Century Gothic, Trebuchet MS, sans-serif", size=11),
        connector=dict(line=dict(color="#B8B8B8", width=1, dash="dot")),
        increasing=dict(marker=dict(color=M2P_BLUE)),
        decreasing=dict(marker=dict(color=RED)),
        totals=dict(marker=dict(color=M2P_GREEN)),
        hovertemplate="%{x}<br>%{y:,.0f} m²<extra></extra>",
    ))
    fig_wf.update_layout(
        title="From available land to effective PV area (m²)",
        yaxis=dict(title="Area (m²)", tickformat=",.0f"),
        showlegend=False,
        plot_bgcolor="white", paper_bgcolor="white",
        font=dict(family="Century Gothic, Trebuchet MS, sans-serif", size=11),
        margin=dict(l=60, r=30, t=60, b=80),
        height=460,
    )
    st.plotly_chart(fig_wf, use_container_width=True)

    total_reduction_pct = (
        (1 - effective_m2 / gate1_available_m2) * 100 if gate1_available_m2 else 0.0
    )
    c1, c2, c3 = st.columns(3)
    c1.metric("Starting available area", f"{gate1_available_m2:,.0f} m²",
              help=f"= {gate1_available_m2/10_000:,.1f} ha")
    c2.metric("Total reduction", f"{total_reduction_pct:.1f}%",
              delta=f"−{gate1_available_m2 - effective_m2:,.0f} m²", delta_color="inverse")
    c3.metric("Effective PV area", f"{effective_m2:,.0f} m²",
              help=f"= {effective_m2/10_000:,.2f} ha  ·  {indicative_mw:,.1f} MW")

    with st.expander("View Calculation Methodology — area reduction"):
        st.markdown(f"""
| Step | Area (m²) | Running total (m²) | Basis |
|---|---|---|---|
| Available area from Gate 1 | — | **{gate1_available_m2:,.0f}** | Total site less built, protected and committed land |
| Less glare zones | −{glare_zones_m2:,.0f} | {gate1_available_m2 - glare_zones_m2:,.0f} | {"Formal glare study" if glare_study_done else "**Study not completed — figure provisional**"} |
| Less OLS restrictions | −{ols_eliminated_m2:,.0f} | {gate1_available_m2 - glare_zones_m2 - ols_eliminated_m2:,.0f} | {"OLS mapping" if ols_study_done else "**Mapping not completed — figure provisional**"} |
| Less runway protection zones | −{rpz_eliminated_m2:,.0f} | **{gate2_cleared_m2:,.0f}** | Absolute exclusion, Gate 2 output |
| Less layout and setbacks | −{layout_loss_m2:,.0f} | **{effective_m2:,.0f}** | {usable_fraction:.0%} usable fraction applied at Gate 6 |

**Effective PV area** = {gate2_cleared_m2:,.0f} m² × {usable_fraction:.2f} = **{effective_m2:,.0f} m²**

**Indicative capacity** = {effective_m2:,.0f} ÷ ({m2_per_mw_lookup[install_type]:,.0f} × {terrain_area_mult:.2f}) = **{indicative_mw:,.2f} MW**

Overall, **{total_reduction_pct:.1f}%** of the land available at Gate 1 is unavailable for panels
once aviation constraints and layout efficiency are applied.
""")

    # ══ 7.2 SATELLITE MAP ═══════════════════════════════════
    st.markdown('<hr class="m2p-divider">', unsafe_allow_html=True)
    st.markdown("#### Footprint at geographic scale")

    if not FOLIUM_AVAILABLE:
        st.warning("Interactive mapping is unavailable. Add `folium` and `streamlit-folium` "
                   "to requirements.txt to enable this view.")
    elif effective_m2 <= 0:
        st.info("No effective PV area to plot. Complete Gates 1, 2 and 6 first.")
    else:
        mc1, mc2, mc3 = st.columns([2, 2, 3])
        with mc1:
            map_lat = st.number_input("Site latitude", -90.0, 90.0, 50.0379, 0.0001,
                                       format="%0.4f")
        with mc2:
            map_lon = st.number_input("Site longitude", -180.0, 180.0, 8.5622, 0.0001,
                                       format="%0.4f")
        with mc3:
            st.caption("Default is Frankfurt Airport (EDDF). Change the coordinates to centre "
                       "the map on the site being assessed.")

        # Square footprint of equal area, drawn at true geographic scale
        side_m = math.sqrt(effective_m2)
        m_per_deg_lat = 111_320.0
        m_per_deg_lon = 111_320.0 * max(math.cos(math.radians(map_lat)), 1e-6)
        d_lat = (side_m / 2) / m_per_deg_lat
        d_lon = (side_m / 2) / m_per_deg_lon

        polygon = [
            [map_lat - d_lat, map_lon - d_lon],
            [map_lat - d_lat, map_lon + d_lon],
            [map_lat + d_lat, map_lon + d_lon],
            [map_lat + d_lat, map_lon - d_lon],
        ]

        fmap = folium.Map(location=[map_lat, map_lon], zoom_start=13, tiles=None,
                          control_scale=True)

        # Esri World Imagery — licensed for this use, so it is the default layer.
        folium.TileLayer(
            tiles=("https://server.arcgisonline.com/ArcGIS/rest/services/"
                   "World_Imagery/MapServer/tile/{z}/{y}/{x}"),
            attr="Esri, Maxar, Earthstar Geographics",
            name="Satellite (Esri)",
            overlay=False,
            control=True,
        ).add_to(fmap)

        # Google Satellite. NOTE: accessing Google's tile servers directly falls
        # outside the Google Maps Platform Terms of Service. Retained as a
        # selectable layer, but Esri is the default for external-facing work.
        folium.TileLayer(
            tiles="https://mt1.google.com/vt/lyrs=s&x={x}&y={y}&z={z}",
            attr="Google",
            name="Satellite (Google — see licensing note)",
            overlay=False,
            control=True,
        ).add_to(fmap)

        folium.TileLayer("OpenStreetMap", name="Street map", overlay=False,
                         control=True).add_to(fmap)

        tooltip_html = (
            f"<b>Indicative PV footprint</b><br>"
            f"Effective area: {effective_m2:,.0f} m² ({effective_m2/10_000:,.2f} ha)<br>"
            f"Indicative capacity: {indicative_mw:,.1f} MW<br>"
            f"Installation: {install_type}<br>"
            f"Square side: {side_m:,.0f} m"
        )

        folium.Polygon(
            locations=polygon,
            color="#2CA02C",
            weight=2,
            fill=True,
            fill_color="#2CA02C",
            fill_opacity=0.35,
            tooltip=folium.Tooltip(tooltip_html, sticky=True),
            popup=folium.Popup(tooltip_html, max_width=320),
        ).add_to(fmap)

        folium.LayerControl(collapsed=True).add_to(fmap)

        st_folium(fmap, width=None, height=520, returned_objects=[])

        st.caption(
            f"The green square has an area of {effective_m2:,.0f} m², equal to the calculated "
            f"effective PV area, drawn at true geographic scale ({side_m:,.0f} m per side). "
            f"It shows the size of the array, not a proposed layout — actual siting depends on "
            f"the glare study, obstacle limitation surfaces and operational constraints."
        )


# ── EXPORT TAB ─────────────────────────────────────────────
with tabs[7]:
    st.markdown('<div class="section-label">Export</div>', unsafe_allow_html=True)
    st.markdown("Download the underlying data for further analysis, or a formatted report "
                "for the client file.")

    run_date = date.today().strftime("%d %B %Y")

    # ── Build the summary KPI frame ─────────────────────────
    summary_rows = [
        ("Project name",                    project_name,                      ""),
        ("Report date",                     run_date,                          ""),
        ("Tool version",                    "3.0",                             ""),
        ("", "", ""),
        ("SITE", "", ""),
        ("Airport development type",        airport_type,                      ""),
        ("Region",                          airport_region,                    ""),
        ("Peak sun hours",                  round(peak_sun, 2),                "h/day"),
        ("Total site area",                 round(total_site_m2, 0),           "m²"),
        ("Built / operational",             built_pct,                         "%"),
        ("Protected ecology area",          round(green_protected_m2, 0),      "m²"),
        ("Committed to other development",  round(committed_dev_m2, 0),        "m²"),
        ("Gate 1 — available area",         round(gate1_available_m2, 0),       "m²"),
        ("Gate 1 — available area (ha)",    round(gate1_available_m2/10_000, 2), "ha"),
        ("", "", ""),
        ("AVIATION CONSTRAINTS", "", ""),
        ("Glare study completed",           "Yes" if glare_study_done else "No", ""),
        ("Area lost to glare",              round(glare_zones_m2, 0),          "m²"),
        ("OLS mapping completed",           "Yes" if ols_study_done else "No", ""),
        ("Area lost to OLS",                round(ols_eliminated_m2, 0),       "m²"),
        ("Area within RPZ",                 round(rpz_eliminated_m2, 0),       "m²"),
        ("Gate 2 — aviation-cleared area",  round(gate2_cleared_m2, 0),        "m²"),
        ("Gate 2 — aviation-cleared area (ha)", round(gate2_cleared_m2/10_000, 2), "ha"),
        ("Reduction from Gate 1",           round(reduction_pct, 1),           "%"),
        ("", "", ""),
        ("GRID", "", ""),
        ("Annual electricity demand",       round(annual_demand_gwh, 1),       "GWh/yr"),
        ("Grid connection headroom",        round(grid_headroom_mw, 1),        "MW"),
        ("Self-consumption",                self_consume_pct,                  "%"),
        ("Export permitted",                "Yes" if export_allowed else "No", ""),
        ("Battery storage included",        "Yes" if add_bess else "No",       ""),
        ("Battery capacity",                round(bess_mwh, 1),                "MWh"),
        ("", "", ""),
        ("SIZING", "", ""),
        ("Installation type",               install_type,                      ""),
        ("Usable fraction",                 round(usable_fraction * 100, 0),   "%"),
        ("Effective PV area",               round(effective_m2, 0),            "m²"),
        ("Effective PV area (ha)",          round(effective_m2/10_000, 2),     "ha"),
        ("Area per MW",                     round(m2_per_mw, 0),               "m²/MW"),
        ("Indicative capacity",             round(indicative_mw, 2),           "MW"),
        ("Capacity override applied",       "Yes" if capacity_override else "No", ""),
        ("", "", ""),
        ("THERMAL PERFORMANCE", "", ""),
        ("Peak ambient temperature",        amb_temp_peak,                     "degC"),
        ("Annual temperature swing",        amb_temp_swing,                    "degC"),
        ("Temperature coefficient",         round(temp_coeff * 100, 2),        "%/degC"),
        ("NOCT",                            noct,                              "degC"),
        ("Peak cell temperature",           round(peak_cell_temp, 1),          "degC"),
        ("Standard annual yield",           round(std_annual_kwh / 1e6, 2),    "GWh"),
        ("Efficiency lost to heat",         round(heat_loss_pct, 2),           "%"),
        ("Worst month",                     worst_month,                       ""),
        ("Worst month loss",                round(worst_derate_pct, 2),        "%"),
        ("Bifacial modules",                "Yes" if bifacial else "No",       ""),
        ("Bifacial gain",                   round(bifacial_gain_pct, 2),       "%"),
        ("Net versus standard yield",       round(net_vs_std_pct, 2),          "%"),
        ("Year 1 generation",               round(base_kwh / 1e6, 2),          "GWh"),
        ("CO2 avoided",                     round(annual_co2_t, 0),            "t/yr"),
        ("", "", ""),
        ("FINANCIALS", "", ""),
        ("Commercial model",                commercial_model.split("(")[0].strip(), ""),
        ("Solar CAPEX rate",                capex_mw,                          "EUR/MW"),
        ("Solar CAPEX",                     round(capex_solar, 0),             "EUR"),
        ("BESS CAPEX",                      round(capex_bess, 0),              "EUR"),
        ("Land cost",                       round(capex_land, 0),              "EUR"),
        ("Total CAPEX",                     round(total_capex, 0),             "EUR"),
        ("Debt ratio",                      round(debt_pct * 100, 0),          "%"),
        ("Debt principal",                  round(debt_principal, 0),          "EUR"),
        ("Equity required",                 round(equity_invest, 0),           "EUR"),
        ("Interest rate",                   round(interest_pct * 100, 2),      "%"),
        ("Loan term",                       loan_years,                        "years"),
        ("Annual debt service",             round(annual_debt_svc, 0),         "EUR/yr"),
        ("Tax rate",                        round(tax_pct * 100, 0),           "%"),
        ("PPA price",                       ppa_price,                         "EUR/kWh"),
        ("After-tax levered IRR",           irr_str,                           ""),
        ("Equity payback",                  pb_str,                            ""),
    ]
    df_summary = pd.DataFrame(summary_rows, columns=["Metric", "Value", "Unit"])

    # ── EXCEL EXPORT ────────────────────────────────────────
    def build_excel() -> bytes:
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
            df_summary.to_excel(writer, sheet_name="Summary KPIs", index=False)

            monthly_out = df_month.copy()
            monthly_out["Loss to heat (%)"] = (
                (1 - monthly_out["Derate factor"]) * 100).round(2)
            monthly_out.to_excel(writer, sheet_name="Monthly Generation", index=False)

            cash_out = df.copy()
            cash_out.to_excel(writer, sheet_name="Projected Cash Flow", index=False)

            assumptions_df = pd.DataFrame({
                "Assumption / limitation": [
                    "Cell temperature uses the NOCT model at 800 W/m2 reference irradiance.",
                    "Monthly ambient temperature is a sinusoid peaking in July.",
                    "Bifacial gain and thermal derating are independent effects, reported separately.",
                    "Soiling losses are NOT modelled (15-30% in Gulf conditions without cleaning).",
                    "Battery revenue is a user-input annual value per MWh, not a dispatch model.",
                    "CO2 factor is the EU average of 0.233 kg/kWh.",
                    "Land is excluded from the depreciable base.",
                    "IRR and payback are pre-feasibility estimates only.",
                    "Glare and OLS assessments must be completed before any layout is committed.",
                ]
            })
            assumptions_df.to_excel(writer, sheet_name="Assumptions", index=False)

            # Column widths
            for sheet_name, frame in [("Summary KPIs", df_summary),
                                      ("Monthly Generation", monthly_out),
                                      ("Projected Cash Flow", cash_out),
                                      ("Assumptions", assumptions_df)]:
                ws = writer.sheets[sheet_name]
                for i, col in enumerate(frame.columns, start=1):
                    longest = max([len(str(col))] +
                                  [len(str(v)) for v in frame[col].head(80)])
                    ws.column_dimensions[ws.cell(row=1, column=i).column_letter].width = \
                        min(max(longest + 2, 12), 60)
        return buffer.getvalue()

    # ── WORD REPORT ─────────────────────────────────────────
    def build_docx() -> bytes:
        doc = Document()

        # Base style
        style = doc.styles["Normal"]
        style.font.name = "Century Gothic"
        style.font.size = Pt(10)

        NAVY = DocxRGB(0x06, 0x34, 0x43)
        BLUE = DocxRGB(0x25, 0x81, 0xC4)

        def heading(text, size=14, colour=NAVY, space_before=14, space_after=6):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after = Pt(space_after)
            r = p.add_run(text)
            r.font.name = "Century Gothic"; r.font.size = Pt(size)
            r.font.bold = True; r.font.color.rgb = colour
            return p

        def body(text, size=10, italic=False, colour=None):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(text)
            r.font.name = "Century Gothic"; r.font.size = Pt(size); r.font.italic = italic
            if colour: r.font.color.rgb = colour
            return p

        def kv_table(rows, widths=(3.4, 2.2, 1.0)):
            t = doc.add_table(rows=0, cols=len(widths))
            t.style = "Light Grid Accent 1"
            t.alignment = WD_TABLE_ALIGNMENT.LEFT
            for row in rows:
                cells = t.add_row().cells
                for i, val in enumerate(row):
                    cells[i].width = Inches(widths[i])
                    para = cells[i].paragraphs[0]
                    run = para.add_run(str(val))
                    run.font.name = "Century Gothic"; run.font.size = Pt(9)
                    if i == 0: run.font.bold = True
            return t

        # ── TITLE PAGE ──
        for _ in range(5): doc.add_paragraph()
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run("Airport Solar Feasibility Assessment")
        r.font.name = "Century Gothic"; r.font.size = Pt(26); r.font.bold = True
        r.font.color.rgb = NAVY

        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(24)
        r = p.add_run(project_name)
        r.font.name = "Century Gothic"; r.font.size = Pt(15); r.font.color.rgb = BLUE

        kv_table([
            ("Prepared by", "M2P Consulting", ""),
            ("Report date", run_date, ""),
            ("Airport type", airport_type, ""),
            ("Region", airport_region.split("·")[0].strip(), ""),
            ("Tool version", "3.0", ""),
            ("Status", "Pre-feasibility — indicative only", ""),
        ], widths=(2.0, 4.0, 0.6))

        doc.add_paragraph()
        body("This report is produced for M2P internal use. All figures are pre-feasibility "
             "estimates and are not suitable for investment decisions without a completed "
             "glare study, obstacle limitation surface mapping, a distribution network "
             "operator grid enquiry, and jurisdiction-specific tax and legal advice.",
             size=9, italic=True)

        doc.add_page_break()

        # ── 1. EXECUTIVE SUMMARY ──
        heading("1.  Executive summary", size=16, space_before=0)
        body(f"The site provides {gate1_available_m2:,.0f} m² ({gate1_available_m2/10_000:,.1f} ha) "
             f"of land after built, protected and "
             f"committed areas are excluded. Aviation safety constraints reduce this by "
             f"{reduction_pct:.0f}% to {gate2_cleared_m2:,.0f} m² of cleared area. Applying a "
             f"{usable_fraction:.0%} usable fraction for setbacks and access gives an indicative "
             f"capacity of {indicative_mw:,.1f} MW using {install_type.lower()}.")
        body(f"Year 1 generation is estimated at {base_kwh/1e6:,.1f} GWh, avoiding approximately "
             f"{annual_co2_t:,.0f} tonnes of CO2 annually. Total capital cost is "
             f"EUR{total_capex:,.0f}, of which EUR{equity_invest:,.0f} is equity at the assumed "
             f"{debt_pct:.0%} debt ratio. The after-tax levered IRR is {irr_str} with equity "
             f"payback at {pb_str.lower()}.")
        if heat_loss_pct > 8:
            body(f"Thermal derating is material at this location: the array loses "
                 f"{heat_loss_pct:.1f}% of standard yield across the year, peaking at "
                 f"{worst_derate_pct:.1f}% in {worst_month}. Soiling losses of a further "
                 f"15-30% apply in arid conditions and are not included in these figures.")

        heading("Gate 1 — site viability", size=12)
        kv_table([
            ("Total site area", f"{total_site_m2:,.0f}", "m²"),
            ("Less built / operational", f"-{built_m2:,.0f}", "m²"),
            ("Less protected ecology", f"-{green_protected_m2:,.0f}", "m²"),
            ("Less committed development", f"-{committed_dev_m2:,.0f}", "m²"),
            ("Available area", f"{gate1_available_m2:,.0f}", "m²"),
            ("Terrain", terrain.split("(")[0].strip(), ""),
            ("Land ownership", land_ownership.split("-")[0].strip(), ""),
        ])

        heading("Gate 2 — aviation safety constraints", size=12)
        kv_table([
            ("Glare study completed", "Yes" if glare_study_done else "No — area may be overstated", ""),
            ("Area lost to glare", f"-{glare_zones_m2:,.0f}", "m²"),
            ("OLS mapping completed", "Yes" if ols_study_done else "No — restrictions unknown", ""),
            ("Area lost to OLS", f"-{ols_eliminated_m2:,.0f}", "m²"),
            ("Area within runway protection zone", f"-{rpz_eliminated_m2:,.0f}", "m²"),
            ("Aviation-cleared area", f"{gate2_cleared_m2:,.0f}", "m²"),
            ("Aviation-cleared area (ha)", f"{gate2_cleared_m2/10_000:,.2f}", "ha"),
            ("Reduction from Gate 1", f"{reduction_pct:.0f}", "%"),
        ])

        doc.add_page_break()

        # ── 2. SIZING AND FINANCIALS ──
        heading("2.  Indicative sizing and financials", size=16, space_before=0)

        heading("Sizing", size=12)
        kv_table([
            ("Installation type", install_type, ""),
            ("Effective PV area after usable fraction", f"{effective_m2:,.0f}", "m²"),
            ("Area per MW", f"{m2_per_mw:,.0f}", "m²/MW"),
            ("Indicative capacity", f"{indicative_mw:,.2f}", "MW"),
            ("Year 1 generation", f"{base_kwh/1e6:,.2f}", "GWh"),
            ("CO2 avoided annually", f"{annual_co2_t:,.0f}", "t"),
        ])

        heading("Capital cost", size=12)
        kv_table([
            ("Solar", f"EUR{capex_solar:,.0f}", ""),
            ("Battery storage", f"EUR{capex_bess:,.0f}", ""),
            ("Land acquisition", f"EUR{capex_land:,.0f}", ""),
            ("Total CAPEX", f"EUR{total_capex:,.0f}", ""),
            ("Debt", f"EUR{debt_principal:,.0f}", f"{debt_pct:.0%}"),
            ("Equity required", f"EUR{equity_invest:,.0f}", f"{1-debt_pct:.0%}"),
        ])

        heading("Returns", size=12)
        kv_table([
            ("After-tax levered IRR", irr_str, ""),
            ("Equity payback", pb_str, ""),
            ("Project life", f"{project_life}", "years"),
            ("PPA / avoided cost", f"EUR{ppa_price:.3f}", "/kWh"),
            ("Commercial model", commercial_model.split("(")[0].strip(), ""),
        ])

        heading("Thermal performance", size=12)
        kv_table([
            ("Peak ambient temperature", f"{amb_temp_peak}", "degC"),
            ("Peak cell temperature", f"{peak_cell_temp:.0f}", "degC"),
            ("Temperature coefficient", f"{temp_coeff*100:.2f}", "%/degC"),
            ("Standard annual yield", f"{std_annual_kwh/1e6:,.2f}", "GWh"),
            ("Efficiency lost to heat", f"{heat_loss_pct:.1f}", "%"),
            ("Worst month", f"{worst_month} ({worst_derate_pct:.1f}% loss)", ""),
            ("Bifacial gain", f"{bifacial_gain_pct:.1f}", "%"),
            ("Net versus standard yield", f"{net_vs_std_pct:+.1f}", "%"),
        ])

        doc.add_page_break()

        # ── 3. CALCULATION BASIS ──
        heading("3.  Calculation basis", size=16, space_before=0)
        body("Every figure in this report is traceable to the steps below.")

        heading("3.1  Site area to installed capacity", size=12)
        body(f"Effective area = cleared area x usable fraction = "
             f"{gate2_cleared_m2:,.0f} m² x {usable_fraction:.2f} = {effective_m2:,.0f} m²")
        body(f"Capacity = effective area / (area per MW x terrain multiplier) = "
             f"{effective_m2:,.0f} m² / ({m2_per_mw_lookup[install_type]:,.0f} m²/MW x "
             f"{terrain_area_mult:.2f}) = {indicative_mw:,.2f} MW")
        if capacity_override:
            body("A client-specified capacity override is applied, so the area-derived figure "
                 "above does not constrain the result.", italic=True)

        heading("3.2  Cell temperature and thermal derating", size=12)
        body(f"Cell temperature = ambient + (NOCT - 20) / 800 x 800 W/m2. "
             f"With NOCT at {noct} degC this gives a constant rise of {noct-20:.0f} degC "
             f"above ambient, so peak cell temperature is {peak_cell_temp:.0f} degC.")
        body(f"Derate factor = 1 + ({temp_coeff*100:.2f}% x (cell temperature - 25 degC)). "
             f"Averaged across the year this gives a {heat_loss_pct:.1f}% reduction against "
             f"standard yield.")

        heading("3.3  Bifacial gain", size=12)
        if bifacial:
            body(f"Rear-side gain = albedo x bifaciality x view factor = {albedo:.2f} x "
                 f"{bifaciality:.2f} x {view_factor:.2f} = +{albedo*bifaciality*view_factor*100:.1f}%. "
                 f"The view factor is set by installation type; elevated ground-mount admits "
                 f"substantially more reflected light than a flush-mounted roof.")
            body("Bifacial gain and thermal derating are independent physical effects. Bifacial "
                 "adds rear-side irradiance; it does not recover heat losses.", italic=True)
        else:
            body("Bifacial modules are not selected, so no rear-side gain is applied.")

        heading("3.4  Capital cost", size=12)
        body(f"Solar = {indicative_mw:,.2f} MW x EUR{capex_mw:,.0f}/MW x {terrain_capex_mult:.2f} "
             f"terrain multiplier = EUR{capex_solar:,.0f}")
        if add_bess:
            body(f"Battery = {bess_mwh:,.1f} MWh x EUR{bess_capex_mwh:,.0f}/MWh = EUR{capex_bess:,.0f}")
        body(f"Total CAPEX = EUR{total_capex:,.0f}, split EUR{debt_principal:,.0f} debt and "
             f"EUR{equity_invest:,.0f} equity.")

        heading("3.5  Annual cash flow", size=12)
        body(f"Net cash flow = revenue - operating cost - lifecycle cost - debt service - tax.")
        body(f"Year 1 revenue = {base_kwh/1e6:,.2f} GWh x EUR{ppa_price:.3f}/kWh x "
             f"{self_consume_pct}% self-consumption. PPA escalates at {ppa_escalation:.1%} annually.")
        body(f"Operating cost = EUR{capex_solar * opex_pct:,.0f} in year 1, inflating at "
             f"{opex_inflation:.1%}. Inverter replacement of EUR{indicative_mw*60000:,.0f} "
             f"falls in year 15.")
        body(f"Debt service = EUR{annual_debt_svc:,.0f} per year for {loan_years} years at "
             f"{interest_pct:.2%}. Straight-line depreciation of EUR{annual_dep:,.0f} per year "
             f"applies to the {project_life}-year life, with land excluded from the "
             f"depreciable base.")

        doc.add_page_break()

        # ── 4. ASSUMPTIONS AND LIMITATIONS ──
        heading("4.  Assumptions and limitations", size=16, space_before=0)
        for item in [
            "Cell temperature uses the NOCT model at 800 W/m2 reference irradiance. Actual cell "
            "temperature varies with wind speed, mounting configuration and rear ventilation.",
            "Monthly ambient temperature is modelled as a sinusoid peaking in July. Replace with "
            "local monthly means for a site-specific assessment.",
            "Soiling losses are not modelled. In Gulf conditions dust reduces output by a further "
            "15-30% without a scheduled cleaning programme, which at an airport requires "
            "security-cleared crews operating in restricted zones.",
            "Battery revenue is a user-entered annual value per MWh, not a dispatch model.",
            "CO2 avoided uses the EU average grid factor of 0.233 kg/kWh. Apply a local factor "
            "outside the EU.",
            "IRR and payback are pre-feasibility estimates. A full financial model requires "
            "audited demand data, a formal grid study, and jurisdiction-specific tax advice.",
            "Glare and obstacle limitation surface assessments must be completed before any "
            "layout is committed. Where these are outstanding, the cleared area above is a "
            "maximum rather than a confirmed figure.",
        ]:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(item)
            r.font.name = "Century Gothic"; r.font.size = Pt(9.5)

        doc.add_paragraph()
        body("M2P Consulting  |  Airport Planning and Design  |  Internal use only",
             size=8, italic=True)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    # ── DOWNLOAD BUTTONS ────────────────────────────────────
    st.markdown('<hr class="m2p-divider">', unsafe_allow_html=True)
    col_a, col_b = st.columns(2)

    with col_a:
        st.markdown("**Editable data**")
        st.caption("Four sheets: Summary KPIs, Monthly Generation, Projected Cash Flow, Assumptions.")
        try:
            st.download_button(
                label="Download data (.xlsx)",
                data=build_excel(),
                file_name=f"M2P_Solar_Feasibility_{project_name.replace(' ', '_')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
            )
        except Exception as exc:
            st.error(f"Excel export unavailable: {exc}. Confirm openpyxl is installed.")

    with col_b:
        st.markdown("**Formatted report**")
        st.caption("Executive summary, sizing and financials, calculation basis, limitations.")
        if DOCX_AVAILABLE:
            try:
                st.download_button(
                    label="Download Full Report (.docx)",
                    data=build_docx(),
                    file_name="M2P_Solar_Feasibility_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as exc:
                st.error(f"Report generation failed: {exc}")
        else:
            st.warning("python-docx is not installed. Add `python-docx` to requirements.txt "
                       "to enable the Word report.")

    # ── ON-SCREEN SUMMARY ───────────────────────────────────
    st.markdown('<hr class="m2p-divider">', unsafe_allow_html=True)
    st.markdown('<div class="section-label">Summary preview</div>', unsafe_allow_html=True)
    st.dataframe(
        df_summary[df_summary["Metric"] != ""],
        use_container_width=True, hide_index=True, height=420,
    )

# ── FOOTER ─────────────────────────────────────────────────
st.markdown("""
<div class="disclaimer">
  <strong>Disclaimer:</strong> This tool produces pre-feasibility estimates for M2P internal use only.
  All outputs are indicative and require a completed glare study, OLS mapping, DNO grid connection enquiry,
  and jurisdiction-specific tax, legal and engineering advice before being presented to clients or used for
  investment decisions. Source: M2P Airport Solar Feasibility Tool v3.0, August 2026.
</div>
""", unsafe_allow_html=True)
